from __future__ import annotations

import csv
import json
from pathlib import Path


class Exporter:
    def __init__(self, output_dir: Path):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def per_image_text(self, result: dict) -> Path:
        path = self.output_dir / f"{Path(result['name']).stem}.txt"
        path.write_text(result["text"], encoding="utf-8")
        return path

    def export(self, results: list[dict], fmt: str) -> Path:
        stem = self.output_dir / "OCR_Result"
        if fmt == "txt":
            path = stem.with_suffix(".txt")
            path.write_text("\n\n".join(f"===== {r['name']} =====\n{r['text']}" for r in results), encoding="utf-8")
        elif fmt == "json":
            path = stem.with_suffix(".json")
            path.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
        elif fmt == "csv":
            path = stem.with_suffix(".csv")
            with path.open("w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f); writer.writerow(["file", "text", "confidence"])
                for r in results:
                    for line in r["lines"]: writer.writerow([r["name"], line["text"], line["confidence"]])
        elif fmt == "docx":
            from docx import Document
            path = stem.with_suffix(".docx"); doc = Document()
            for r in results: doc.add_heading(r["name"], 2); doc.add_paragraph(r["text"])
            doc.save(path)
        elif fmt == "xlsx":
            from openpyxl import Workbook
            path = stem.with_suffix(".xlsx"); wb = Workbook(); ws = wb.active; ws.title = "OCR Results"; ws.append(["File", "Text", "Confidence"])
            for r in results:
                for line in r["lines"]: ws.append([r["name"], line["text"], line["confidence"]])
            wb.save(path)
        elif fmt == "pdf":
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            path = stem.with_suffix(".pdf"); styles = getSampleStyleSheet(); story = []
            for r in results: story.extend([Paragraph(r["name"], styles["Heading2"]), Paragraph(r["text"].replace("\n", "<br/>"), styles["BodyText"]), Spacer(1, 12)])
            SimpleDocTemplate(str(path), pagesize=A4).build(story)
        else: raise ValueError(f"Unsupported format: {fmt}")
        return path
